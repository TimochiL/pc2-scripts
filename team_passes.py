import random
import string
import sys

from docx import Document


def main():
    team_count = int(sys.argv[1])
    passwords = list()
    passwords_doc = Document()
    team_index = 1
    while team_index <= team_count:
        password = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
        passwords.append(password)
        add_team_slip(passwords_doc, str(team_index), password)
        team_index += 1
    export_text_list(passwords)
    passwords_doc.save("passwords.docx")


def add_team_slip(document, team_number, password):
    document.add_heading("Novice Team" + team_number)
    document.add_paragraph("Welcome to the 2024 Clements Computer Science Competition!")
    team_name = 'team' + team_number
    document.add_paragraph("PC2 Username: " + team_name)
    document.add_paragraph("PC2 Password: " + password)


def export_text_list(passwords):
    with open("passwords.txt", 'w') as password_file:
        password_file.write('\n'.join(passwords))


if __name__ == '__main__':
    main()
